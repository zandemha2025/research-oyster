"""Build report.docx — the results-first written report, as a real Word document.

Mirrors the markdown report's structure (POV → executive answer → the numbers → themes →
recommendations → method → confidence/limitations → sources) but styled as a shareable .docx
a consultant would send. Built directly from the synthesis dict (not by parsing markdown), so
headings, the metrics tables, and the [n] ledger get proper Word styles.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from reporting import theme

_INK = RGBColor(*theme.rgb(theme.INK))
_SUBTLE = RGBColor(*theme.rgb(theme.SUBTLE))
_ACCENT = RGBColor(*theme.rgb(theme.ACCENT))


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = theme.FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _INK


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = theme.FONT
    run.font.color.rgb = _ACCENT if level == 1 else _INK
    run.font.size = Pt(15 if level == 1 else 12.5)


def _para(doc: Document, text: str, *, color: RGBColor | None = None, size: float = 10.5,
          italic: bool = False, bold: bool = False) -> None:
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = theme.FONT
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.bold = bold
    run.font.color.rgb = color or _INK


def _metric_table(doc: Document, table: dict[str, Any]) -> None:
    from reporting.tables import normalize_table

    norm = normalize_table(table)
    if not norm["body"]:
        return
    _heading(doc, norm["title"] + (f"  ({norm['unit']})" if norm["unit"] else ""), level=2)
    header, body = norm["header"], norm["body"]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    for cell, label in zip(t.rows[0].cells, header):
        run = cell.paragraphs[0].add_run(str(label))
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.name = theme.FONT
    for row in body:
        cells = t.add_row().cells
        for cell, v in zip(cells, row):
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(9.5)
            run.font.name = theme.FONT
    if table.get("note"):
        _para(doc, table["note"], color=_SUBTLE, size=9, italic=True)


def build_docx(job: dict[str, Any], synthesis: dict[str, Any], generated_at: str,
               charts: list[dict[str, Any]] | None, path: Path, *,
               themes: list[dict[str, Any]] | None = None,
               biblio: list[dict[str, Any]] | None = None, title: str = "") -> Path:
    doc = Document()
    _style(doc)

    _heading(doc, title or job.get("brief") or "Research report", level=1)
    _para(doc, f"Research Oyster · {generated_at} · job #{job.get('id')}", color=_SUBTLE, size=9)

    if synthesis.get("point_of_view"):
        _heading(doc, "Our point of view", level=2)
        _para(doc, synthesis["point_of_view"], size=12, bold=True)

    if synthesis.get("key_takeaways"):
        _heading(doc, "At a glance", level=2)
        for t in synthesis["key_takeaways"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(t))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.name = theme.FONT

    _heading(doc, "Executive answer", level=2)
    _para(doc, synthesis.get("executive_answer") or "—")

    tables = synthesis.get("metrics_tables") or []
    if tables:
        _heading(doc, "The numbers", level=1)
        for table in tables:
            _metric_table(doc, table)
        # Embed the rendered charts under the numbers, if we have them.
        for ch in charts or []:
            png = ch.get("png")
            if png and Path(png).exists():
                try:
                    doc.add_picture(str(png))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

    themes = themes if themes is not None else (synthesis.get("themes") or [])
    if themes:
        _heading(doc, "What players are saying", level=1)
        for th in themes:
            _heading(doc, th.get("title") or "Theme", level=2)
            if th.get("insight"):
                _para(doc, th["insight"])
            for c in th.get("citations") or []:
                quote, src, n = c.get("quote"), (c.get("source") or c.get("url")), c.get("n")
                if quote:
                    line = f'"{quote}"' + (f" — {src}" if src else "") + (f" [{n}]" if n else "")
                elif src:
                    line = src + (f" [{n}]" if n else "")
                else:
                    continue
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.name = theme.FONT

    if synthesis.get("tensions"):
        _heading(doc, "Tensions and disagreement", level=2)
        _para(doc, synthesis["tensions"])
    if synthesis.get("sentiment"):
        _heading(doc, "Sentiment", level=2)
        _para(doc, synthesis["sentiment"])

    recs = synthesis.get("recommendations") or []
    if recs:
        _heading(doc, "Recommendations", level=1)
        for rec in recs:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(str(rec))
            run.font.size = Pt(10.5)
            run.font.name = theme.FONT

    if synthesis.get("method"):
        _heading(doc, "Method — how the numbers were made", level=2)
        _para(doc, synthesis["method"], size=10)

    if synthesis.get("confidence") or synthesis.get("limitations"):
        _heading(doc, "Confidence and limitations", level=2)
        _para(doc, synthesis.get("confidence") or "")
        if synthesis.get("limitations"):
            _para(doc, synthesis["limitations"], color=_SUBTLE, size=10)

    sources = biblio if biblio is not None else (synthesis.get("numbered_sources") or [])
    if sources:
        _heading(doc, "Sources", level=1)
        for s in sources:
            p = doc.add_paragraph()
            run = p.add_run(f"[{s.get('n')}] {s.get('label') or ''}".rstrip())
            run.font.size = Pt(9.5)
            run.font.name = theme.FONT
            run.font.color.rgb = _INK
            if s.get("url"):
                run2 = p.add_run(f"  {s['url']}")
                run2.font.size = Pt(9)
                run2.font.name = theme.FONT
                run2.font.color.rgb = _SUBTLE

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path
